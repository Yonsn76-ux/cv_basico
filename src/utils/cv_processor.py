# -*- coding: utf-8 -*-
"""
Procesador de CVs simplificado para clasificación por profesiones
"""

import os
import re
import PyPDF2
import pytesseract
from docx import Document
import cv2
import numpy as np
from PIL import Image
import pandas as pd

class CVProcessor:
    """Procesador simplificado de CVs"""

    DEFAULT_SKILL_KEYWORDS = [
        # Tecnología
        'python', 'java', 'javascript', 'sql', 'html', 'css', 'react', 'angular',
        'node.js', 'php', 'c++', 'c#', '.net', 'spring', 'django', 'flask',
        'git', 'docker', 'kubernetes', 'aws', 'azure', 'linux', 'windows',

        # Data Science
        'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas',
        'numpy', 'scikit-learn', 'tableau', 'power bi', 'excel', 'r',
        'statistics', 'data analysis', 'big data', 'hadoop', 'spark',

        # Marketing
        'marketing digital', 'seo', 'sem', 'google ads', 'facebook ads',
        'social media', 'content marketing', 'email marketing', 'analytics',
        'photoshop', 'illustrator', 'canva', 'hootsuite',

        # Diseño
        'diseño gráfico', 'ui/ux', 'figma', 'sketch', 'adobe creative',
        'after effects', 'premiere', 'indesign', 'branding', 'tipografía',

        # Ventas
        'ventas', 'sales', 'crm', 'salesforce', 'negociación', 'prospección',
        'atención al cliente', 'customer service', 'retail',

        # Administración
        'administración', 'gestión', 'management', 'liderazgo', 'proyectos',
        'planificación', 'presupuestos', 'finanzas', 'contabilidad', 'rrhh',

        # Agricultura
        'agricultura', 'agronomía', 'cultivos', 'riego', 'fertilizantes',
        'pesticidas', 'maquinaria agrícola', 'ganadería', 'veterinaria',
        'producción agrícola', 'agropecuario', 'campo'
    ]

    def __init__(self, skill_keywords=None):
        """Inicializa el procesador.

        Parameters
        ----------
        skill_keywords : list[str] or None
            Lista de palabras clave para detectar habilidades en el texto. Si se
            pasa ``None`` se usa la lista por defecto. Una lista vacía desactiva
            la detección de habilidades.
        """
        self.supported_formats = ['.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png', '.bmp', '.tiff']
        # Permitir personalizar las keywords de habilidades o deshabilitarlas
        if skill_keywords is None:
            self.skill_keywords = self.DEFAULT_SKILL_KEYWORDS
        else:
            self.skill_keywords = skill_keywords
    
    def extract_text_from_file(self, file_path):
        """Extrae texto de un archivo según su formato"""
        try:
            file_ext = os.path.splitext(file_path.lower())[1]
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_word(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                return self._extract_from_image(file_path)
            else:
                return ""
        except Exception as e:
            print(f"Error procesando {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, pdf_path):
        """Extrae texto de PDF"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error en PDF {pdf_path}: {e}")
        return text
    
    def _extract_from_word(self, word_path):
        """Extrae texto de documento Word"""
        text = ""
        try:
            doc = Document(word_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"Error en Word {word_path}: {e}")
        return text
    
    def _extract_from_image(self, image_path):
        """Extrae texto de imagen usando OCR"""
        try:
            # Preprocesar imagen
            image = cv2.imread(image_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Mejorar contraste
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            gray = clahe.apply(gray)
            
            # OCR
            config = '--lang spa --oem 3 --psm 6'
            text = pytesseract.image_to_string(gray, config=config)
            return text
        except Exception as e:
            print(f"Error en imagen {image_path}: {e}")
            # Fallback con PIL
            try:
                img = Image.open(image_path)
                text = pytesseract.image_to_string(img, lang='spa')
                return text
            except:
                return ""
    
    def clean_text(self, text):
        """Limpia y normaliza el texto"""
        if not text:
            return ""
        
        # Convertir a minúsculas
        text = text.lower()
        
        # Remover caracteres especiales pero mantener espacios y puntos
        text = re.sub(r'[^\w\s\.\@\-\+\(\),;:]', ' ', text)
        
        # Normalizar espacios
        text = re.sub(r'\s+', ' ', text)
        
        # Remover líneas muy cortas
        lines = text.split('\n')
        clean_lines = [line.strip() for line in lines if len(line.strip()) > 3]
        
        return ' '.join(clean_lines).strip()
    
    def extract_features(self, text):
        """Extrae características específicas del CV"""
        features = {
            'emails': [],
            'phones': [],
            'experience_years': [],
            'education_keywords': [],
            'skills': [],
            'text_length': len(text),
            'word_count': len(text.split())
        }
        
        # Extraer emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        features['emails'] = re.findall(email_pattern, text)
        
        # Extraer teléfonos
        phone_pattern = r'[\+]?[\d\s\-\(\)]{8,15}'
        features['phones'] = re.findall(phone_pattern, text)
        
        # Buscar años de experiencia
        exp_patterns = [
            r'(\d+)\s*años?\s*de\s*experiencia',
            r'experiencia\s*de\s*(\d+)\s*años?',
            r'(\d+)\s*years?\s*of\s*experience'
        ]
        for pattern in exp_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            features['experience_years'].extend(matches)
        
        # Keywords de educación
        education_keywords = [
            'universidad', 'licenciatura', 'ingeniería', 'maestría', 'doctorado',
            'técnico', 'certificación', 'diplomado', 'bachillerato', 'carrera',
            'university', 'bachelor', 'master', 'phd', 'degree'
        ]
        for keyword in education_keywords:
            if keyword in text:
                features['education_keywords'].append(keyword)
        
        # Keywords de habilidades
        for skill in self.skill_keywords:
            if skill in text:
                features['skills'].append(skill)
        
        return features
    
    def process_cv_folder(self, folder_path, profession_name):
        """Procesa todos los CVs de una carpeta para una profesión específica"""
        results = []
        
        if not os.path.exists(folder_path):
            print(f"La carpeta {folder_path} no existe")
            return results
        
        files = [f for f in os.listdir(folder_path) 
                if os.path.splitext(f.lower())[1] in self.supported_formats]
        
        print(f"Procesando {len(files)} archivos para la profesión: {profession_name}")
        
        for file_name in files:
            file_path = os.path.join(folder_path, file_name)
            
            # Extraer texto
            raw_text = self.extract_text_from_file(file_path)
            clean_text = self.clean_text(raw_text)
            
            if clean_text:
                # Extraer características
                features = self.extract_features(clean_text)
                
                result = {
                    'file_name': file_name,
                    'profession': profession_name,
                    'text': clean_text,
                    'features': features,
                    'status': 'success'
                }
            else:
                result = {
                    'file_name': file_name,
                    'profession': profession_name,
                    'text': '',
                    'features': {},
                    'status': 'failed'
                }
            
            results.append(result)
            print(f"  ✓ {file_name} - {result['status']}")
        
        return results
    
    def is_supported_file(self, file_path):
        """Verifica si el archivo tiene un formato soportado"""
        ext = os.path.splitext(file_path.lower())[1]
        return ext in self.supported_formats
